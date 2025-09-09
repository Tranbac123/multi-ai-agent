"""Document processor for ingestion pipeline."""

import asyncio
import hashlib
import mimetypes
import time
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import structlog
import aiofiles
import PyPDF2
import docx
import openpyxl
from PIL import Image
import json

logger = structlog.get_logger(__name__)


class DocumentProcessor:
    """Processes documents for ingestion pipeline."""
    
    def __init__(self, max_file_size: int = 100 * 1024 * 1024):  # 100MB
        self.max_file_size = max_file_size
        self.supported_types = {
            'text/plain': self._process_text,
            'text/markdown': self._process_text,
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_xlsx,
            'image/jpeg': self._process_image,
            'image/png': self._process_image,
            'application/json': self._process_json,
        }
    
    async def process_document(
        self,
        file_path: str,
        tenant_id: str,
        user_id: Optional[str] = None,
        permissions: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Process a document and return metadata and chunks."""
        try:
            # Validate file
            file_info = await self._validate_file(file_path)
            if not file_info:
                raise ValueError("Invalid file")
            
            # Generate document hash for idempotency
            doc_hash = await self._generate_file_hash(file_path)
            
            # Check if document already processed
            if await self._is_document_processed(doc_hash, tenant_id):
                logger.info("Document already processed", doc_hash=doc_hash, tenant_id=tenant_id)
                return await self._get_existing_document(doc_hash, tenant_id)
            
            # Process document based on type
            mime_type = file_info['mime_type']
            if mime_type not in self.supported_types:
                raise ValueError(f"Unsupported file type: {mime_type}")
            
            processor = self.supported_types[mime_type]
            content, metadata = await processor(file_path)
            
            # Chunk content
            chunks = await self._chunk_content(content, metadata)
            
            # Create document metadata
            document_metadata = {
                'doc_id': str(uuid.uuid4()),
                'tenant_id': tenant_id,
                'user_id': user_id,
                'doc_hash': doc_hash,
                'file_path': file_path,
                'file_name': Path(file_path).name,
                'file_size': file_info['size'],
                'mime_type': mime_type,
                'created_at': time.time(),
                'permissions': permissions or {},
                'metadata': metadata,
                'chunks': chunks,
                'status': 'processed'
            }
            
            # Store document metadata
            await self._store_document_metadata(document_metadata)
            
            logger.info(
                "Document processed successfully",
                doc_id=document_metadata['doc_id'],
                tenant_id=tenant_id,
                chunks_count=len(chunks)
            )
            
            return document_metadata
            
        except Exception as e:
            logger.error(
                "Document processing failed",
                error=str(e),
                file_path=file_path,
                tenant_id=tenant_id
            )
            raise
    
    async def _validate_file(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Validate file and return file info."""
        try:
            path = Path(file_path)
            if not path.exists():
                return None
            
            # Check file size
            size = path.stat().st_size
            if size > self.max_file_size:
                raise ValueError(f"File too large: {size} bytes")
            
            # Detect MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            if not mime_type:
                # Try to detect from file extension
                mime_type = self._detect_mime_type_from_extension(path.suffix)
            
            return {
                'size': size,
                'mime_type': mime_type,
                'extension': path.suffix
            }
            
        except Exception as e:
            logger.error("File validation failed", error=str(e), file_path=file_path)
            return None
    
    async def _generate_file_hash(self, file_path: str) -> str:
        """Generate SHA-256 hash of file content."""
        hash_sha256 = hashlib.sha256()
        async with aiofiles.open(file_path, 'rb') as f:
            while chunk := await f.read(8192):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    async def _is_document_processed(self, doc_hash: str, tenant_id: str) -> bool:
        """Check if document is already processed."""
        # This would typically check a database
        # For now, return False
        return False
    
    async def _get_existing_document(self, doc_hash: str, tenant_id: str) -> Dict[str, Any]:
        """Get existing document metadata."""
        # This would typically fetch from database
        # For now, return empty dict
        return {}
    
    async def _store_document_metadata(self, metadata: Dict[str, Any]) -> None:
        """Store document metadata."""
        # This would typically store in database
        logger.info("Document metadata stored", doc_id=metadata['doc_id'])
    
    def _detect_mime_type_from_extension(self, extension: str) -> str:
        """Detect MIME type from file extension."""
        extension_map = {
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.json': 'application/json',
        }
        return extension_map.get(extension.lower(), 'application/octet-stream')
    
    async def _process_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process text file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        
        metadata = {
            'type': 'text',
            'encoding': 'utf-8',
            'line_count': len(content.splitlines()),
            'word_count': len(content.split())
        }
        
        return content, metadata
    
    async def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process PDF file."""
        content = ""
        metadata = {
            'type': 'pdf',
            'page_count': 0,
            'has_images': False,
            'has_tables': False
        }
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata['page_count'] = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    content += page_text + "\n"
                
                # Check for images and tables (simplified)
                if '/XObject' in pdf_reader.pages[0].get('/Resources', {}):
                    metadata['has_images'] = True
                
        except Exception as e:
            logger.error("PDF processing failed", error=str(e), file_path=file_path)
            raise
        
        return content, metadata
    
    async def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX file."""
        content = ""
        metadata = {
            'type': 'docx',
            'paragraph_count': 0,
            'has_tables': False,
            'has_images': False
        }
        
        try:
            doc = docx.Document(file_path)
            metadata['paragraph_count'] = len(doc.paragraphs)
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Check for tables
            if doc.tables:
                metadata['has_tables'] = True
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            content += cell.text + " "
                        content += "\n"
            
            # Check for images (simplified)
            if any(rel.target_ref for rel in doc.part.rels.values() if 'image' in rel.target_ref):
                metadata['has_images'] = True
                
        except Exception as e:
            logger.error("DOCX processing failed", error=str(e), file_path=file_path)
            raise
        
        return content, metadata
    
    async def _process_xlsx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process XLSX file."""
        content = ""
        metadata = {
            'type': 'xlsx',
            'sheet_count': 0,
            'has_formulas': False
        }
        
        try:
            workbook = openpyxl.load_workbook(file_path)
            metadata['sheet_count'] = len(workbook.sheetnames)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content += f"Sheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join(str(cell) for cell in row if cell is not None)
                    content += row_text + "\n"
                
        except Exception as e:
            logger.error("XLSX processing failed", error=str(e), file_path=file_path)
            raise
        
        return content, metadata
    
    async def _process_image(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process image file."""
        try:
            with Image.open(file_path) as img:
                metadata = {
                    'type': 'image',
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode,
                    'has_transparency': img.mode in ('RGBA', 'LA', 'P')
                }
                
                # For images, we'll store basic metadata as content
                content = f"Image: {Path(file_path).name}\n"
                content += f"Dimensions: {img.width}x{img.height}\n"
                content += f"Format: {img.format}\n"
                content += f"Mode: {img.mode}\n"
                
        except Exception as e:
            logger.error("Image processing failed", error=str(e), file_path=file_path)
            raise
        
        return content, metadata
    
    async def _process_json(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process JSON file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        
        try:
            json_data = json.loads(content)
            metadata = {
                'type': 'json',
                'is_valid_json': True,
                'structure': self._analyze_json_structure(json_data)
            }
        except json.JSONDecodeError:
            metadata = {
                'type': 'json',
                'is_valid_json': False,
                'structure': {}
            }
        
        return content, metadata
    
    def _analyze_json_structure(self, data: Any) -> Dict[str, Any]:
        """Analyze JSON structure."""
        if isinstance(data, dict):
            return {
                'type': 'object',
                'keys': list(data.keys()),
                'key_count': len(data)
            }
        elif isinstance(data, list):
            return {
                'type': 'array',
                'length': len(data),
                'item_types': list(set(type(item).__name__ for item in data))
            }
        else:
            return {
                'type': type(data).__name__
            }
    
    async def _chunk_content(self, content: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Chunk content into smaller pieces."""
        chunks = []
        chunk_size = 1000  # characters
        overlap = 100  # characters
        
        if len(content) <= chunk_size:
            # Single chunk
            chunks.append({
                'chunk_id': str(uuid.uuid4()),
                'content': content,
                'start_index': 0,
                'end_index': len(content),
                'metadata': metadata
            })
        else:
            # Multiple chunks with overlap
            start = 0
            chunk_index = 0
            
            while start < len(content):
                end = min(start + chunk_size, len(content))
                chunk_content = content[start:end]
                
                chunks.append({
                    'chunk_id': str(uuid.uuid4()),
                    'content': chunk_content,
                    'start_index': start,
                    'end_index': end,
                    'chunk_index': chunk_index,
                    'metadata': metadata
                })
                
                start = end - overlap
                chunk_index += 1
        
        return chunks